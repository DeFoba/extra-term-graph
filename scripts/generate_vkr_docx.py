import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Импортируем контент по главам
from vkr_content import vkr_title, vkr_intro, vkr_chapter1, vkr_chapter2, vkr_chapter3, vkr_conclusion, vkr_references, vkr_appendix

def create_document():
    doc = Document()

    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Times New Roman'
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.bold = True
        if level == 1:
            hs.font.size = Pt(16)
        elif level == 2:
            hs.font.size = Pt(14)
        else:
            hs.font.size = Pt(14)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

    # Поля страниц
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    return doc

def add_paragraph(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, font_size=14, first_line_indent=Cm(1.25)):
    p = doc.add_paragraph()
    p.alignment = align
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    return p

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
    doc.add_paragraph()
    return table

def render_content(doc, content_blocks):
    """
    Отрисовывает блоки контента в документ.
    Блок может быть:
    {"type": "p", "text": "...", "bold": False, "align": ...}
    {"type": "h1", "text": "..."}
    {"type": "h2", "text": "..."}
    {"type": "h3", "text": "..."}
    {"type": "table", "headers": [...], "rows": [...]}
    {"type": "img", "path": "...", "caption": "..."}
    {"type": "pb"} - page break
    """
    for block in content_blocks:
        if block['type'] == 'p':
            bold = block.get('bold', False)
            align_val = block.get('align', WD_ALIGN_PARAGRAPH.JUSTIFY)
            if align_val == "CENTER":
                align_val = WD_ALIGN_PARAGRAPH.CENTER
            indent = block.get('indent', Cm(1.25))
            fs = block.get('font_size', 14)
            add_paragraph(doc, block['text'], bold=bold, align=align_val, first_line_indent=indent, font_size=fs)
        elif block['type'] == 'h1':
            add_heading(doc, block['text'], level=1)
        elif block['type'] == 'h2':
            add_heading(doc, block['text'], level=2)
        elif block['type'] == 'h3':
            add_heading(doc, block['text'], level=3)
        elif block['type'] == 'table':
            add_table(doc, block['headers'], block['rows'])
        elif block['type'] == 'img':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(block['path'], width=Cm(16))
            if 'caption' in block:
                caption = doc.add_paragraph(block['caption'])
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in caption.runs:
                    r.font.size = Pt(12)
                    r.font.italic = True
        elif block['type'] == 'pb':
            doc.add_page_break()

def generate():
    doc = create_document()
    
    print("Рендеринг Титульного листа...")
    render_content(doc, vkr_title.get_content())

    print("Рендеринг Оглавления...")
    add_heading(doc, "ОГЛАВЛЕНИЕ", level=1)
    toc_items = [
        ("ВВЕДЕНИЕ", ""),
        ("1. ТЕОРЕТИКО-МЕТОДОЛОГИЧЕСКИЕ ОСНОВЫ ИНТЕЛЛЕКТУАЛЬНОГО АНАЛИЗА НАУЧНЫХ ТЕКСТОВ", ""),
        ("   1.1 Проблема поиска и систематизации научных публикаций", ""),
        ("   1.2 Обзор и сравнительный анализ алгоритмов извлечения ключевых слов", ""),
        ("   1.3 Автоматическое реферирование научных текстов", ""),
        ("   1.4 Графовое представление знаний и СУБД Neo4j", ""),
        ("   1.5 Выводы по главе", ""),
        ("2. ПРОЕКТИРОВАНИЕ И АРХИТЕКТУРА СИСТЕМЫ ИНТЕЛЛЕКТУАЛЬНОГО АНАЛИЗА", ""),
        ("   2.1 Формирование и анализ функциональных и нефункциональных требований", ""),
        ("   2.2 Обоснование выбора стека технологий", ""),
        ("   2.3 Разработка архитектуры и конвейера обработки данных (Pipeline)", ""),
        ("   2.4 Проектирование графовой модели данных", ""),
        ("   2.5 Выбор математических метрик для оценки качества работы", ""),
        ("   2.6 Выводы по главе", ""),
        ("3. ПРАКТИЧЕСКАЯ РЕАЛИЗАЦИЯ И ЭКСПЕРИМЕНТАЛЬНОЕ ТЕСТИРОВАНИЕ", ""),
        ("   3.1 Программная реализация модуля сборки и очистки корпуса", ""),
        ("   3.2 Программная реализация подсистемы извлечения терминов", ""),
        ("   3.3 Программная реализация модуля автоматического реферирования", ""),
        ("   3.4 Программная реализация графовой интеграции и веб-приложения", ""),
        ("   3.5 Экспериментальное тестирование и анализ результатов", ""),
        ("   3.6 Выводы по главе", ""),
        ("ЗАКЛЮЧЕНИЕ", ""),
        ("СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ И ЛИТЕРАТУРЫ", ""),
        ("ПРИЛОЖЕНИЯ", ""),
    ]
    for title, _ in toc_items:
        add_paragraph(doc, title, first_line_indent=None)
    doc.add_page_break()
    
    print("Рендеринг Введения...")
    render_content(doc, vkr_intro.get_content())
    
    print("Рендеринг Главы 1...")
    render_content(doc, vkr_chapter1.get_content())
    
    print("Рендеринг Главы 2...")
    render_content(doc, vkr_chapter2.get_content())
    
    print("Рендеринг Главы 3...")
    render_content(doc, vkr_chapter3.get_content())
    
    print("Рендеринг Заключения...")
    render_content(doc, vkr_conclusion.get_content())
    
    print("Рендеринг Списка литературы...")
    render_content(doc, vkr_references.get_content())

    print("Рендеринг Приложений...")
    render_content(doc, vkr_appendix.get_content())
    
    out_path = os.path.join(os.path.dirname(__file__), "..", "docs", "ВКР_2026.docx")
    out_path = os.path.abspath(out_path)
    doc.save(out_path)
    print(f"Готово! Документ сохранен: {out_path}")

if __name__ == "__main__":
    generate()
